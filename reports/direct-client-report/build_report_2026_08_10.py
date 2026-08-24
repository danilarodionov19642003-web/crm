from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mentori/mentori-crm")
OUT_DIR = ROOT / "reports/direct-client-report"
OUT = OUT_DIR / "Отчет_Яндекс_Директ_Спектр_Металла_1-10_августа_2026.docx"
ASSET = Path("/Users/mentori/Спектр металл/assets/direct")
MENTORI_LOGO = Path("/Users/mentori/mentra-site/assets/mentori-technologies-wordmark-transparent-light.png")

RED = "C82020"
DARK = "17191F"
CHARCOAL = "292C33"
MID = "626773"
LIGHT = "F3F4F6"
PALE_RED = "FBEAEA"
PALE_GREEN = "EAF5EF"
GREEN = "2E7D50"
WHITE = "FFFFFF"
LINE = "D9DCE2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches, indent_dxa=120):
    widths = [int(v * 1440) for v in widths_inches]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=10.3, bold=False, color=DARK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, before=0, after=6, line=1.12, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep


def add_text(doc, text="", size=10.3, bold=False, color=DARK, after=6,
             before=0, align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, keep=keep)
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=3, line=1.08)
    set_font(p.add_run(text), size=9.8, color=DARK)
    return p


def create_numbering(doc, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    lvl.append(start_node)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text_node = OxmlElement("w:lvlText")
    text_node.set(qn("w:val"), "%1.")
    lvl.append(text_node)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, title, body, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    style_paragraph(p, after=5, line=1.12)
    set_font(p.add_run(title + ". "), size=10.1, bold=True, color=DARK)
    set_font(p.add_run(body), size=10.1, color=DARK)
    return p


def add_callout(doc, label, text, fill=LIGHT, label_color=RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell,
                    top={"val": "single", "sz": 4, "color": fill},
                    bottom={"val": "single", "sz": 4, "color": fill},
                    left={"val": "single", "sz": 22, "color": label_color},
                    right={"val": "single", "sz": 4, "color": fill})
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.15)
    set_font(p.add_run(label + "  "), size=10.2, bold=True, color=label_color)
    set_font(p.add_run(text), size=10.2, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_geometry(table, [6.5 / len(metrics)] * len(metrics))
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, DARK)
        set_cell_border(cell,
                        top={"val": "single", "sz": 4, "color": DARK},
                        bottom={"val": "single", "sz": 4, "color": DARK},
                        left={"val": "single", "sz": 4, "color": CHARCOAL},
                        right={"val": "single", "sz": 4, "color": CHARCOAL})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=1)
        set_font(p.add_run(value), size=18, bold=True, color=WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(label), size=8.4, color="D4D6DB")
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_section_gap(doc, points=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 0.7


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def make_campaign_table(doc):
    data = [
        ["Кампания", "Показы", "Переходы", "CTR", "CPC", "Расход", "Действия"],
        ["Профнастил", "4 699", "507", "10,79%", "12,40 ₽", "6 284,38 ₽", "8"],
        ["Металлочерепица", "1 115", "142", "12,74%", "24,07 ₽", "3 417,27 ₽", "5"],
        ["Итого", "5 814", "649", "11,16%", "14,95 ₽", "9 701,65 ₽", "13"],
    ]
    table = doc.add_table(rows=len(data), cols=7)
    widths = [1.38, .66, .77, .58, .70, 1.05, 1.36]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "D7D9DE" if c_idx == 6 else DARK)
                color, bold = (DARK if c_idx == 6 else WHITE), True
            elif r_idx == len(data) - 1:
                set_cell_shading(cell, LIGHT)
                color, bold = DARK, True
            else:
                set_cell_shading(cell, WHITE)
                color, bold = DARK, c_idx == 0
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 6) else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.7, bold=bold, color=color)
    return table


def add_conversion_table(doc):
    rows = [
        ["Заявка из корзины", "1 - отправлена форма заказа"],
        ["Заявка через форму", "1 - отправлена форма обратной связи"],
        ["Клик по телефону", "10 - посетитель открыл набор номера"],
        ["Переход в WhatsApp", "1 - посетитель открыл диалог с компанией"],
        ["Итого в Метрике", "13 целевых действий, включая 2 отправленные заявки"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2.05, 4.45])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, LIGHT if r_idx == len(rows) - 1 else WHITE)
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.8,
                     bold=(r_idx == len(rows) - 1 or c_idx == 0), color=DARK)
    return table


def add_traffic_source_table(doc):
    rows = [
        ["Рекламное направление", "Визиты", "Пользователи / цели"],
        ["Профнастил", "318", "252 / 8"],
        ["Металлочерепица", "103", "72 / 5"],
        ["Итого из рекламы", "421", "320 / 13"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_geometry(table, [2.35, 1.25, 2.65], indent_dxa=0)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, LIGHT if r_idx == 0 else (LIGHT if r_idx == len(rows) - 1 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.9,
                     bold=(r_idx == 0 or r_idx == len(rows) - 1 or c_idx == 0),
                     color=DARK)
    return table


def add_bad_queries_table(doc):
    rows = [
        ["Тип спроса", "Примеры из отчёта", "Действие"],
        ["Другие города", "Новокузнецк, Тюмень, Лиски", "исключить города"],
        ["Конкуренты и адреса", "Сайдинг Инвест, ИП Игнатенко, Ю3, 33 Северная", "исключить бренды"],
        ["Другие товары", "краска для пола, металлолом, резиновый профнастил", "исключить тематику"],
        ["Полезные вопросы", "размеры, расчёт, выбор профиля, крыша и забор", "перенести в SEO"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_geometry(table, [1.62, 3.25, 1.63])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.5, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_seo_table(doc):
    rows = [
        ["Тема", "Подтверждение спроса", "Что создать"],
        ["Обшивка дома профлистом", "15 показов / 10 переходов по одной формулировке", "Гид по выбору профиля для фасада"],
        ["Забор из профнастила", "54 показа / 8 переходов", "Посадочная страница и расчёт забора"],
        ["Расчёт материала", "Переходы по калькулятору, площади, ширине и длине", "Калькулятор и инструкция по расчёту"],
        ["Цены и размеры листов", "Регулярные запросы по 2, 3 и 6-метровым листам", "Страницы цен за лист и м²"],
        ["Металлочерепица", "Запросы по Монтеррей, цветам, размерам и выбору", "Гид по профилям и кровельному комплекту"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_geometry(table, [1.55, 2.15, 2.80])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.0)
            set_font(p.add_run(value), size=7.9, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_group_results_table(doc):
    rows = [
        ["Рекламная группа", "Переходы", "CTR", "Цена перехода"],
        ["Профнастил: купить и цена", "183", "11,72%", "13,96 ₽"],
        ["Профнастил: цвета и размеры", "163", "9,85%", "12,55 ₽"],
        ["Металлочерепица: крыша и кровля", "102", "12,78%", "21,84 ₽"],
        ["Профнастил: для крыши", "80", "9,71%", "10,05 ₽"],
        ["Профнастил: для забора", "71", "12,54%", "9,80 ₽"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [2.80, 1.05, 1.00, 1.65])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.6, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_acceleration_table(doc):
    rows = [
        ["Что меняем", "Что получает посетитель"],
        ["Мгновенное переключение категорий", "Переходы без повторной загрузки всего каталога"],
        ["Оптимизация шрифтов и анимаций", "Более быстрое появление страницы на телефоне"],
        ["Облегчение фильтров и выбора цвета", "Меньше задержек на старых iPhone и Safari"],
        ["Сохранение отдельных SEO-адресов", "Реклама и поисковая индексация продолжают работать"],
        ["Повторные замеры и тестирование", "Проверенный результат на компьютере и телефоне"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2.85, 3.65])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.7, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_image_grid(doc):
    items = [
        (ASSET / "profnastil-fence-16x9.png", "Готовый забор"),
        (ASSET / "profnastil-roof-16x9.png", "Готовая кровля"),
        (ASSET / "profnastil-colors-4x3.png", "Профили и цвета"),
        (ASSET / "metallocherepitsa-fan-1x1.png", "Линейка цветов"),
        (ASSET / "metallocherepitsa-house-4x3.png", "Кровля, графит"),
        (ASSET / "metallocherepitsa-roof-16x9.png", "Кровля, бордо"),
    ]
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [2.166, 2.167, 2.167])
    for idx, (path, caption) in enumerate(items):
        cell = table.cell(idx // 3, idx % 3)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, WHITE)
        set_cell_border(cell,
                        top={"val": "single", "sz": 7, "color": LINE},
                        bottom={"val": "single", "sz": 7, "color": LINE},
                        left={"val": "single", "sz": 7, "color": LINE},
                        right={"val": "single", "sz": 7, "color": LINE})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=2)
        if path.exists():
            shape = p.add_run().add_picture(str(path), width=Inches(1.82), height=Inches(1.05))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(caption), size=8.3, bold=True, color=DARK)
    return table


def add_budget_table(doc):
    rows = [
        ["Этап", "Бюджет", "Как распределяем", "Цель этапа"],
        ["1. Контрольный тест", "10 000 ₽", "Фактически использовано 9 701,65 ₽", "Гипотезы проверены, спрос и рабочие направления определены"],
        ["2. Масштабирование", "19 690 ₽", "Ориентир: 12 000 ₽ на профнастил и 7 690 ₽ на металлочерепицу", "Усиливаем связки с качественными обращениями"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1.52, .78, 2.0, 2.20])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "D7D9DE" if r_idx == 0 and c_idx == 3 else (DARK if r_idx == 0 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 5, "color": LINE},
                            bottom={"val": "single", "sz": 5, "color": LINE},
                            left={"val": "single", "sz": 5, "color": LINE},
                            right={"val": "single", "sz": 5, "color": LINE})
            p = cell.paragraphs[0]
            if r_idx == 0 and c_idx == 3:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.08)
            text_color = DARK if r_idx == 0 and c_idx == 3 else (WHITE if r_idx == 0 else DARK)
            set_font(p.add_run(value), size=8.5, bold=(r_idx == 0 or c_idx in (0, 1)), color=text_color)
    return table


def add_header_footer(doc):
    for section in doc.sections:
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(hp, after=0)
        set_font(hp.add_run("СПЕКТР МЕТАЛЛА  |  ЯНДЕКС.ДИРЕКТ  |  MENTORI"), size=7.8, bold=True, color=MID)
        fp = section.footer.paragraphs[0]
        style_paragraph(fp, after=0)
        set_font(fp.add_run("Отчёт за 1-10 августа 2026  •  metallomsk.ru  •  mentori.tech"), size=7.8, color=MID)
        fp.add_run(" " * 8)
        run = fp.add_run("Страница ")
        set_font(run, size=7.8, color=MID)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, (size, color, before, after) in enumerate(((18, RED, 16, 8), (13.5, DARK, 11, 5), (11.3, CHARCOAL, 8, 4)), start=1):
        style = styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.8)

    # Page 1: cover and executive summary.
    add_text(doc, "СПЕКТР МЕТАЛЛА", size=10, bold=True, color=RED, after=28)
    add_text(doc, "Результаты рекламы\nи точки роста", size=28, bold=True, color=DARK, after=8, keep=True)
    add_text(doc, "Профнастил и металлочерепица  •  Омск и Омская область", size=12.8, bold=True, color=MID, after=3)
    add_text(doc, "Срез за 29 июля - 5 августа 2026 года", size=10, color=MID, after=3)
    add_text(doc, "Подготовлено MENTORI  •  mentori.tech", size=9.4, bold=True, color=RED, after=22)
    add_metric_strip(doc, [
        ("3 479", "показов"),
        ("387", "переходов"),
        ("6 756 ₽", "расход с НДС"),
        ("12", "целевых действий"),
    ])
    add_text(doc, "Главный вывод", size=14, bold=True, color=DARK, before=18, after=8)
    add_callout(doc, "Результат", "Обе кампании приводят целевые действия по близкой стоимости. Профнастил обеспечивает основной объём трафика, а металлочерепица заметно улучшила кликабельность к концу периода.", fill=PALE_GREEN, label_color=GREEN)
    add_bullet(doc, "Средняя стоимость перехода составила 14,31 ₽ при установленном пределе 70 ₽.")
    add_bullet(doc, "Зафиксировано 8 кликов по телефону, 2 перехода в WhatsApp и 2 отправленные формы.")
    add_bullet(doc, "4-5 августа: 1 665 показов, 186 переходов и 7 целевых действий.")
    add_bullet(doc, "Расход за период: 5 537,63 ₽ без НДС или 6 755,91 ₽ с НДС 22%.")
    add_bullet(doc, "Нецелевые переходы обнаружены и собраны в отдельный список для дополнительной очистки.")
    add_bullet(doc, "151 формулировка связана с ценами, выбором, размерами, цветами и профилями - это отдельная возможность для SEO-продвижения.")
    add_text(doc, "План бюджета", size=14, bold=True, color=DARK, before=12, after=7)
    add_callout(doc, "30 000 ₽", "10 000 ₽ - контролируемый тест и очистка трафика. Оставшиеся 20 000 ₽ - усиление направлений, которые подтверждаются качественными обращениями.", fill=PALE_RED, label_color=RED)

    # Page 2: campaign performance.
    add_page_break(doc)
    section_one_numbers = create_numbering(doc)
    add_heading(doc, "1. Результаты рекламных кампаний", 1)
    add_text(doc, "Целевое действие - отправка формы или попытка связаться с компанией: клик по телефону либо переход в мессенджер. Это важный шаг к заявке, но ещё не означает состоявшуюся продажу.", size=9.2, color=MID, after=10)
    make_campaign_table(doc)
    add_callout(doc, "Расход с НДС", "Фактическая стоимость рекламы для клиента за период - 6 755,91 ₽. В таблице сохранены данные Директа без НДС, по которым рассчитываются CPC и CPA.", fill=PALE_RED, label_color=RED)
    add_heading(doc, "Что показывают цифры", 2)
    add_number(doc, "Профнастил даёт основной объём", "297 переходов, CTR 11,21% и средняя стоимость перехода 12,90 ₽. Кампания собрала 8 целевых действий.", section_one_numbers)
    add_number(doc, "Металлочерепица набрала темп", "90 переходов привели к 4 целевым действиям. Ориентировочная стоимость одного действия - 426,33 ₽ против 479,04 ₽ у профнастила.", section_one_numbers)
    add_number(doc, "Общий CTR высокий", "11,12% посетителей, увидевших рекламу, перешли на сайт. Средняя стоимость перехода по двум кампаниям - 14,31 ₽.", section_one_numbers)
    add_number(doc, "Последние два дня сильнее среднего", "4-5 августа получено 186 переходов и 7 целевых действий. CPA составил 314,69 ₽ без НДС, что на 31,8% ниже среднего за весь период.", section_one_numbers)
    add_heading(doc, "Сравнение эффективности", 2)
    winners = [
        ("Профнастил", "CPC 12,90 ₽  •  8 действий  •  ориентировочный CPA 479,04 ₽"),
        ("Металлочерепица", "CPC 18,95 ₽  •  4 действия  •  ориентировочный CPA 426,33 ₽"),
        ("Итого", "12 целевых действий  •  ориентировочный CPA 461,47 ₽"),
    ]
    table = doc.add_table(rows=len(winners), cols=2)
    set_table_geometry(table, [2.05, 4.45])
    for idx, (name, detail) in enumerate(winners):
        for col, value in enumerate((name, detail)):
            cell = table.cell(idx, col)
            set_cell_shading(cell, LIGHT if idx % 2 == 0 else WHITE)
            set_cell_border(cell,
                            top={"val": "single", "sz": 3, "color": LINE},
                            bottom={"val": "single", "sz": 3, "color": LINE},
                            left={"val": "single", "sz": 3, "color": LINE},
                            right={"val": "single", "sz": 3, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.9, bold=(col == 0 or idx == 2), color=DARK)
    add_callout(doc, "Важно", "В отчёте не называем двенадцать действий двенадцатью заявками. Отдел продаж должен подтверждать, какие формы, звонки и сообщения стали реальными обращениями и заказами.", fill=LIGHT, label_color=MID)

    # Page 3: conversions and targeting.
    add_page_break(doc)
    add_heading(doc, "2. Контакты и источники трафика", 1)
    add_text(doc, "Директ учитывает действия посетителей на сайте. Для оценки бизнеса эти данные нужно сопоставлять с письмами, звонками и фактическими диалогами менеджера.", after=9)
    add_text(doc, "Какие действия зафиксированы", size=11.3, bold=True, color=DARK, after=5)
    add_conversion_table(doc)
    add_callout(doc, "5 августа", "Метрика увидела 5 действий с рекламы: заявку из корзины, заявку через форму и 3 клика по телефону. В отчёте Директа пока зачтено 4 конверсии; текущий день может доатрибутироваться, а несколько действий одного посетителя не равны нескольким клиентам.", fill=PALE_GREEN, label_color=GREEN)
    add_heading(doc, "Роль автотаргетинга", 2)
    add_traffic_source_table(doc)
    add_bullet(doc, "Автотаргетинг дал 291 из 387 переходов и использовал 77,1% расходов.")
    add_bullet(doc, "Полностью отключать его сейчас невыгодно: он обеспечивает основной объём и уже приводит контактные действия.")
    add_bullet(doc, "Правильная тактика - регулярно чистить поисковые запросы и ограничивать явно нецелевые темы.")
    add_callout(doc, "Статус", "Кампании работают только в поиске Яндекса, по Омску и Омской области, ежедневно с 08:00 до 21:00 по омскому времени.", fill=LIGHT, label_color=MID)

    # Page 4: search-query cleanup.
    add_page_break(doc)
    section_queries = create_numbering(doc)
    add_heading(doc, "3. Качество поисковых запросов", 1)
    add_callout(doc, "Применено 2 августа", "Пакет минус-фраз внесён: 14 добавлений в кампанию профнастила и 13 - в кампанию металлочерепицы. Он закрыл часть нецелевых запросов, но не все формулировки, обнаруженные позднее. Бюджеты, стратегии, география и расписание не менялись.", fill=PALE_GREEN, label_color=GREEN)
    add_bad_queries_table(doc)
    add_heading(doc, "Что делаем с этим трафиком", 2)
    add_number(doc, "Расширяем минус-фразы", "Добавляем фамилии и названия компаний, другие города, адресные запросы и неподходящие товары.", section_queries)
    add_number(doc, "Не блокируем полезный спрос", "Запрос «металлический штакетник Омск» уже дал контактное действие. Его лучше вынести в отдельную посадочную страницу и рекламную группу.", section_queries)
    add_number(doc, "Информационный спрос переносим в SEO", "39 справочных формулировок дали 44 перехода, 321,63 ₽ расходов и 0 целевых действий. На них лучше отвечать статьями, а не оплачивать каждый переход.", section_queries)
    add_heading(doc, "Что уже закрыто текущими минус-фразами", 2)
    add_bullet(doc, "«профнастил оцинковка Ишим» - исключено словом «ишим»;")
    add_bullet(doc, "«Сайдинг Инвест» - исключено отдельной фразой;")
    add_bullet(doc, "«краска для пола» и «профнастил б/у» - закрываются словами «краска» и «бу».")
    add_heading(doc, "Что ещё требуется добавить", 2)
    add_bullet(doc, "«Новокузнецк», «ИП Игнатенко Анатолий Владимирович», «Основа Строй», «печать», «рифлёный»;")
    add_bullet(doc, "Новые запросы 5 августа: «Лемана ПРО», «Леруа Мерлен», «Алейск», «лист металла 1 мм».")
    add_text(doc, "Эти формулировки сейчас отсутствуют в минус-фразах кампании. Их необходимо включить в следующий пакет очистки после согласования.", size=9.4, bold=True, color=RED, before=7, after=0)

    # Page 5: SEO upsell opportunity.
    add_page_break(doc)
    add_heading(doc, "4. Дополнительная возможность роста: SEO", 1)
    add_callout(doc, "Отдельная услуга", "SEO-продвижение не входит в текущий объём работ и ранее не заказывалось. Ниже показан подтверждённый рекламой спрос, который можно проработать в отдельном SEO-проекте.", fill=PALE_RED, label_color=RED)
    add_text(doc, "За период найдено 151 уникальное поисковое выражение о цене, выборе, размерах, расчёте, цветах и профилях. Они дали 224 оплаченных перехода и использовали 2 797,34 ₽ рекламного бюджета без НДС.", after=9)
    add_seo_table(doc)
    add_text(doc, "Категории пересекаются: один запрос может одновременно относиться к цене, размеру и конкретному профилю. Поэтому строки таблицы нельзя складывать между собой.", size=8.5, italic=True, color=MID, before=5, after=8)
    add_heading(doc, "Что предлагается допродать", 2)
    add_bullet(doc, "Сбор и кластеризация полной SEO-семантики по Омску и области.")
    add_bullet(doc, "Посадочные страницы по ценам, профилям, покрытиям и применению.")
    add_bullet(doc, "Статьи про выбор для крыши и забора, размеры и расчёт материала.")
    add_bullet(doc, "Техническая SEO-настройка, перелинковка, индексация и контроль обращений.")
    add_callout(doc, "Приоритет", "Начать с цен, выбора для крыши и забора, размеров и калькулятора. Реклама уже подтвердила интерес к этим темам, но SEO потребует отдельного согласования и бюджета.", fill=PALE_GREEN, label_color=GREEN)

    # Page 6: messages and visual tests.
    add_page_break(doc)
    add_heading(doc, "5. Что тестируем в объявлениях", 1)
    add_text(doc, "Рекламные группы разделены по задаче покупателя. Это позволяет сравнить не только товары, но и аргументы, формулировки и изображения.", after=9)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3.19, 3.31])
    blocks = [
        ("ПРОФНАСТИЛ", [
            "Купить профнастил в Омске",
            "Профнастил для забора",
            "Профлист для крыши",
            "Профили С8-НС44",
            "Цвета RAL и размеры",
        ]),
        ("МЕТАЛЛОЧЕРЕПИЦА", [
            "Цена и покупка в Омске",
            "Металлочерепица для крыши",
            "Кровельный комплект",
            "Супер-Монтеррей",
            "Размеры и выбор цвета",
        ]),
    ]
    for col, (title, lines) in enumerate(blocks):
        cell = table.cell(0, col)
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell,
                        top={"val": "single", "sz": 5, "color": LINE},
                        bottom={"val": "single", "sz": 5, "color": LINE},
                        left={"val": "single", "sz": 5, "color": LINE},
                        right={"val": "single", "sz": 5, "color": LINE})
        p = cell.paragraphs[0]
        style_paragraph(p, after=5)
        set_font(p.add_run(title), size=10.2, bold=True, color=RED)
        for line in lines:
            p = cell.add_paragraph(style="List Bullet")
            style_paragraph(p, after=2, line=1.05)
            set_font(p.add_run(line), size=8.8, color=DARK)
    add_heading(doc, "Какие аргументы и фотографии сравниваем", 2)
    add_bullet(doc, "Цена, изготовление по размерам, выбор RAL, доставка по Омску и области.")
    add_bullet(doc, "Готовый объект против наглядной фотографии отдельного профиля или цветовой линейки.")
    add_image_grid(doc)
    add_text(doc, "Следующий вывод по визуалам делаем не по вкусу, а по фактической цене качественного обращения после накопления данных.", size=9.1, color=MID, before=7, after=0)

    # Page 7: budget and next steps.
    add_page_break(doc)
    section_next = create_numbering(doc)
    add_heading(doc, "6. Бюджет и следующий этап", 1)
    add_text(doc, "Общий план остаётся прежним: сначала ограниченный тест, затем масштабирование только подтверждённых связок.", after=10)
    add_budget_table(doc)
    add_callout(doc, "Текущий остаток", "С НДС из тестовых 10 000 ₽ использовано 6 755,91 ₽; остаётся 3 244,09 ₽.", fill=PALE_RED, label_color=RED)
    add_heading(doc, "Что делаем дальше", 2)
    add_number(doc, "Дочищаем запросы", "вносим новые минус-фразы и ежедневно проверяем, куда уходит бюджет автотаргетинга.", section_next)
    add_number(doc, "Продолжаем обе кампании", "профнастил сохраняет сильный спрос, металлочерепица пока показывает более низкую стоимость контактного действия.", section_next)
    add_number(doc, "Сверяем с продажами", "менеджер помечает качество каждого звонка, сообщения и письма, чтобы оптимизация шла по реальным обращениям.", section_next)
    add_number(doc, "Предлагаем отдельный SEO-этап", "сначала согласовываем объём и стоимость, затем начинаем с цены, выбора для крыши и забора, размеров и калькулятора.", section_next)
    add_number(doc, "После тестовых 10 000 ₽", "распределяем оставшиеся 20 000 ₽ по цене качественной заявки и согласованным приоритетам бизнеса.", section_next)
    add_heading(doc, "Что нужно подтвердить у клиента", 2)
    add_bullet(doc, "Какие профили, покрытия и цвета сейчас важнее по наличию и маржинальности.")
    add_bullet(doc, "Какие обращения считаются качественными и какие завершились продажей.")
    add_bullet(doc, "Нужно ли следующим отдельным направлением запускать евроштакетник.")
    add_bullet(doc, "Готов ли клиент отдельно заказать SEO-продвижение по выявленным темам спроса.")
    add_callout(doc, "Итог", "Реклама уже даёт измеримые контактные действия. Ближайшая задача - уменьшить долю нецелевого спроса и связать статистику Директа с реальными продажами.", fill=PALE_GREEN, label_color=GREEN)
    add_text(doc, "Отчёт подготовлен MENTORI  •  mentori.tech", size=10.2, bold=True, color=RED, before=9, after=0)

    add_header_footer(doc)
    doc.core_properties.title = "Результаты Яндекс.Директа и точки роста - Спектр Металла"
    doc.core_properties.subject = "Профнастил и металлочерепица, 29 июля - 5 августа 2026"
    doc.core_properties.author = "MENTORI"
    doc.core_properties.keywords = "Яндекс.Директ, SEO, профнастил, металлочерепица, Омск, отчёт, MENTORI"
    doc.save(OUT)
    print(OUT)


def build_august():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, (size, color, before, after) in enumerate(
            ((18, RED, 16, 8), (13.5, DARK, 11, 5), (11.3, CHARCOAL, 8, 4)), start=1):
        style = doc.styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.8)

    # Page 1: executive summary.
    add_text(doc, "СПЕКТР МЕТАЛЛА", size=10, bold=True, color=RED, after=28)
    add_text(doc, "Итоги тестового этапа\nи план роста", size=28, bold=True, color=DARK, after=8, keep=True)
    add_text(doc, "Профнастил и металлочерепица  •  Омск и Омская область", size=12.8, bold=True, color=MID, after=3)
    add_text(doc, "Отчёт за 1-10 августа 2026 года", size=10, color=MID, after=3)
    add_text(doc, "Подготовлено MENTORI  •  mentori.tech", size=9.4, bold=True, color=RED, after=22)
    add_metric_strip(doc, [
        ("5 814", "показов"),
        ("649", "переходов"),
        ("9 702 ₽", "расход"),
        ("13", "целевых действий"),
    ])
    add_text(doc, "Главный вывод", size=14, bold=True, color=DARK, before=18, after=8)
    add_callout(doc, "Тест завершён", "Плановый тестовый лимит 10 000 ₽ практически полностью использован. Мы проверили две товарные категории, рекламные формулировки, изображения и реальные поисковые запросы. Полученных данных достаточно, чтобы перейти от широкого теста к управляемому масштабированию.", fill=PALE_GREEN, label_color=GREEN)
    add_bullet(doc, "Средняя стоимость перехода - 14,95 ₽ при установленном ограничении 70 ₽.")
    add_bullet(doc, "Зафиксированы 2 отправленные заявки и 11 попыток связаться по телефону или в WhatsApp.")
    add_bullet(doc, "Профнастил обеспечивает основной объём и обе отправленные заявки; металлочерепица показывает высокий интерес, но более дорогой переход.")
    add_bullet(doc, "В рекламных запросах выявлен устойчивый спрос на расчёты, размеры, цены, выбор материала и применение - это основа для отдельного SEO-этапа.")
    add_text(doc, "Следующий бюджетный этап", size=14, bold=True, color=DARK, before=12, after=7)
    add_callout(doc, "19 690 ₽", "Текущий остаток кабинета на 11 августа. Предлагается направить его на лучшие связки, продолжая контролировать качество запросов и реальные обращения отдела продаж.", fill=PALE_RED, label_color=RED)

    # Page 2: campaign performance.
    add_page_break(doc)
    section_results = create_numbering(doc)
    add_heading(doc, "1. Результаты рекламных кампаний", 1)
    add_text(doc, "Все суммы приведены с учётом НДС. Целевое действие - отправка заявки либо попытка связаться с компанией. Оно показывает интерес, но не равно подтверждённой продаже.", size=9.2, color=MID, after=10)
    make_campaign_table(doc)
    add_heading(doc, "Что показывают цифры", 2)
    add_number(doc, "Профнастил - основное направление", "507 переходов по 12,40 ₽. Кампания дала 8 действий, включая обе отправленные формы заказа и обратной связи.", section_results)
    add_number(doc, "Металлочерепица подтверждает спрос", "142 перехода и 5 контактных действий. CTR 12,74% выше, чем у профнастила, но переход стоит 24,07 ₽, поэтому кампанию нужно масштабировать осторожнее.", section_results)
    add_number(doc, "Общая кликабельность стабильна", "CTR 11,16% означает, что объявления хорошо отвечают поисковому спросу и привлекают внимание целевой аудитории.", section_results)
    add_number(doc, "Тестовый лимит соблюдён", "За закрытый период использовано 9 701,65 ₽. Бюджет не был увеличен автоматически и остался в согласованных рамках.", section_results)
    add_callout(doc, "Вывод", "На следующем этапе профнастил сохраняет приоритет по объёму и подтверждённым заявкам. Металлочерепицу продолжаем параллельно, но усиливаем только те группы и формулировки, которые дают качественные контакты.", fill=PALE_GREEN, label_color=GREEN)

    # Page 3: conversions and site behavior.
    add_page_break(doc)
    add_heading(doc, "2. Обращения и поведение на сайте", 1)
    add_text(doc, "Метрика связывает рекламный переход с действиями посетителя на сайте. Для оценки продаж эти данные необходимо сопоставлять с письмами, звонками и диалогами менеджеров.", after=9)
    add_text(doc, "Какие действия пришли из рекламы", size=11.3, bold=True, color=DARK, after=5)
    add_conversion_table(doc)
    add_heading(doc, "Рекламные визиты по направлениям", 2)
    add_traffic_source_table(doc)
    add_callout(doc, "Всего на сайте", "За 1-10 августа Метрика зарегистрировала 492 визита, 356 пользователей и 797 просмотров страниц. 421 визит пришёл из оплаченной рекламы Яндекса.", fill=LIGHT, label_color=MID)
    add_bullet(doc, "Средняя продолжительность визита - 1 минута 14 секунд.")
    add_bullet(doc, "Средняя глубина просмотра - 1,62 страницы.")
    add_bullet(doc, "Доля быстрых уходов - 24,6%; показатель оставляет резерв для улучшения посадочных страниц и скорости каталога.")
    add_callout(doc, "Важно", "13 действий - это не 13 продаж. Подтверждённый коммерческий результат нужно отмечать со стороны менеджеров: кто дозвонился, какой запрос был целевым и какие обращения завершились заказом.", fill=PALE_RED, label_color=RED)

    # Page 4: completed test and hypotheses.
    add_page_break(doc)
    section_test = create_numbering(doc)
    add_heading(doc, "3. Что проверили за тестовый этап", 1)
    add_text(doc, "Первые 10 000 ₽ использовались не просто на показы рекламы. Их задача - найти рабочие сочетания товара, запроса, объявления и посадочной страницы.", after=9)
    add_group_results_table(doc)
    add_heading(doc, "Что подтвердилось", 2)
    add_number(doc, "Цена и наличие дают основной объём", "Группа «Профнастил: купить и цена» получила 183 перехода - больше остальных отдельных групп.", section_test)
    add_number(doc, "Хорошо работает конкретная задача", "Запросы для забора и крыши дают переходы дешевле среднего по аккаунту: 9,80 ₽ и 10,05 ₽.", section_test)
    add_number(doc, "Цвета и размеры востребованы", "Эта группа получила 163 перехода и подтверждает интерес к подбору RAL, покрытия и длины листа.", section_test)
    add_number(doc, "Металлочерепица сильнее для крыши", "102 из 142 переходов кампании пришлись на группу «Для крыши и кровли».", section_test)
    add_heading(doc, "Что именно сравнивалось", 2)
    add_bullet(doc, "Формулировки про цену, производство, размеры, цвета, крышу и забор.")
    add_bullet(doc, "Товарные фотографии, цветовые линейки и изображения готовых объектов.")
    add_bullet(doc, "Отдельные группы по назначению товара вместо одного общего объявления.")
    add_bullet(doc, "Показы только в поиске Яндекса, по Омску и Омской области, в рабочее время.")
    add_callout(doc, "Решение", "Широкий тест завершён. Далее бюджет распределяется не поровну, а с приоритетом профнастила и наиболее сильных групп металлочерепицы.", fill=PALE_GREEN, label_color=GREEN)

    # Page 5: query quality.
    add_page_break(doc)
    section_queries = create_numbering(doc)
    add_heading(doc, "4. Качество поисковых запросов", 1)
    add_callout(doc, "Очистка выполнена", "2 августа в обе кампании внесён согласованный пакет минус-фраз: 14 добавлений для профнастила и 13 для металлочерепицы. После этого мониторинг продолжился, потому что новые формулировки появляются ежедневно.", fill=PALE_GREEN, label_color=GREEN)
    add_bad_queries_table(doc)
    add_heading(doc, "Что показал мониторинг", 2)
    add_number(doc, "Нецелевой спрос можно дополнительно сократить", "В отчёте появились новые конкуренты, адреса, другие города и товары. Они подготовлены для следующего пакета исключений.", section_queries)
    add_number(doc, "Полезные вопросы нельзя просто блокировать", "Запросы про выбор, размеры и расчёт пока не всегда дают заявку сразу, но показывают реальный интерес будущих покупателей.", section_queries)
    add_number(doc, "Реклама и SEO решают разные задачи", "Коммерческие запросы продолжаем покупать в Директе, а информационный спрос выгоднее постепенно закрывать собственными страницами и статьями.", section_queries)
    add_heading(doc, "Примеры новых исключений", 2)
    add_bullet(doc, "другие города: Новокузнецк, Тюмень, Лиски;")
    add_bullet(doc, "конкуренты и адреса: Сайдинг Инвест, ИП Игнатенко, Ю3, 33 Северная;")
    add_bullet(doc, "неподходящие товары: краска для пола, металлолом, резиновый профнастил.")
    add_text(doc, "Очистка запросов остаётся регулярной частью ведения рекламы: разового списка минус-фраз недостаточно, потому что Яндекс постоянно находит новые варианты формулировок.", size=9.3, bold=True, color=RED, before=8, after=0)

    # Page 6: SEO proposal.
    add_page_break(doc)
    add_heading(doc, "5. Дополнительная возможность роста: SEO", 1)
    add_callout(doc, "Отдельная услуга", "SEO-продвижение не входило в создание сайта и ведение Яндекс.Директа. За тестовый период реклама собрала фактические запросы аудитории - теперь на их основе можно предложить отдельный долгосрочный этап продвижения.", fill=PALE_RED, label_color=RED)
    add_text(doc, "Сейчас каждый переход по таким запросам оплачивается через Директ. Если создать полезные страницы и статьи, часть этой аудитории со временем сможет приходить из обычного поиска без оплаты каждого клика.", after=9)
    add_seo_table(doc)
    add_heading(doc, "Что входит в отдельный SEO-этап", 2)
    add_bullet(doc, "Сбор и кластеризация запросов по Омску и Омской области.")
    add_bullet(doc, "Создание посадочных страниц по товарам, размерам, покрытиям и назначению.")
    add_bullet(doc, "Подготовка статей и инструкций по выбору и расчёту материала.")
    add_bullet(doc, "Техническая оптимизация, внутренняя перелинковка и контроль индексации.")
    add_bullet(doc, "Отслеживание заявок из бесплатного поиска отдельно от рекламы.")
    add_callout(doc, "Экономический смысл", "SEO не заменяет Директ мгновенно. Его задача - в течение нескольких месяцев сформировать собственный поток поискового трафика и постепенно уменьшить зависимость от постоянной оплаты за каждый переход.", fill=PALE_GREEN, label_color=GREEN)

    # Page 7: site acceleration proposal.
    add_page_break(doc)
    add_heading(doc, "6. Отдельный этап развития: ускорение сайта", 1)
    add_callout(doc, "Почему этот этап появился сейчас", "При запуске сайта приоритетом были полнота каталога, корректные цены и варианты товаров, корзина, мобильная версия, аналитика и отдельные страницы для рекламы. После десяти дней реального трафика появились данные о том, как посетители используют большой каталог. Они показали следующий резерв роста - скорость переходов между категориями на телефонах.", fill=PALE_GREEN, label_color=GREEN)
    add_text(doc, "Это не исправление неработающего сайта: функционал работает, а на компьютере техническая оценка скорости составляет 98 из 100. Предлагается отдельная модернизация архитектуры каталога, чтобы улучшить мобильный результат и сделать переключение разделов визуально мгновенным.", after=9)
    add_acceleration_table(doc)
    add_heading(doc, "Что показал технический аудит", 2)
    add_bullet(doc, "Оценка производительности: 98/100 на компьютере и 74/100 на мобильном профиле.")
    add_bullet(doc, "Каждая категория сейчас имеет отдельный адрес для рекламы и SEO, поэтому при переключении браузер заново открывает страницу.")
    add_bullet(doc, "Основной мобильный вес создают шрифты и повторные визуальные эффекты, а не сервер.")
    add_bullet(doc, "Сервер отвечает быстро; ускорение требует именно переработки клиентской части каталога и повторного тестирования устройств.")
    add_callout(doc, "Результат для бизнеса", "Меньше ожидания между разделами, более комфортный просмотр с телефона и снижение риска, что оплаченный рекламный посетитель уйдёт до выбора товара. Работа оформляется как отдельный этап с резервной копией, тестированием и замерами до/после.", fill=PALE_RED, label_color=RED)

    # Page 8: budget and next plan.
    add_page_break(doc)
    section_next = create_numbering(doc)
    add_heading(doc, "7. Бюджет и следующий этап", 1)
    add_text(doc, "Первый этап дал рабочую статистику. Следующая задача - не расширять рекламу хаотично, а направить остаток на подтверждённые направления и продолжить контроль качества обращений.", after=10)
    add_budget_table(doc)
    add_callout(doc, "Баланс на 11 августа", "В кабинете остаётся 19 690 ₽. Текущий день не включён в статистику отчёта, поскольку кампания 11 августа ещё не завершена.", fill=PALE_RED, label_color=RED)
    add_heading(doc, "План действий", 2)
    add_number(doc, "Масштабируем профнастил", "сохраняем основной объём и усиливаем группы «купить и цена», «для забора», «для крыши», «цвета и размеры».", section_next)
    add_number(doc, "Оптимизируем металлочерепицу", "оставляем сильную группу для крыши и кровли, контролируя более высокую стоимость перехода.", section_next)
    add_number(doc, "Продолжаем очистку", "регулярно проверяем поисковые запросы и добавляем новые минус-фразы без изменения согласованной географии.", section_next)
    add_number(doc, "Сверяем рекламу с продажами", "менеджер отмечает качество звонков, сообщений и заявок, чтобы бюджет распределялся по реальным заказам.", section_next)
    add_number(doc, "Отдельно согласовываем рост сайта", "SEO и ускорение каталога оформляются отдельными этапами с понятным объёмом работ и результатом.", section_next)
    add_heading(doc, "Что требуется от клиента", 2)
    add_bullet(doc, "Подтвердить, какие профили и покрытия сейчас приоритетны по наличию и маржинальности.")
    add_bullet(doc, "Передать обратную связь по качеству полученных звонков, сообщений и двух отправленных заявок.")
    add_bullet(doc, "Рассмотреть отдельное предложение по SEO-продвижению выявленных тем.")
    add_bullet(doc, "Рассмотреть отдельный этап ускорения мобильного каталога.")
    add_callout(doc, "Итог", "Тестовый бюджет отработан: спрос подтверждён, сильные направления определены, нецелевые темы выявлены. Следующий этап - масштабирование на основе данных, а не предположений.", fill=PALE_GREEN, label_color=GREEN)
    add_text(doc, "Отчёт подготовлен MENTORI  •  mentori.tech", size=10.2, bold=True, color=RED, before=9, after=0)

    add_header_footer(doc)
    doc.core_properties.title = "Итоги тестового этапа Яндекс.Директа - Спектр Металла"
    doc.core_properties.subject = "Профнастил и металлочерепица, 1-10 августа 2026"
    doc.core_properties.author = "MENTORI"
    doc.core_properties.keywords = "Яндекс.Директ, SEO, ускорение сайта, профнастил, металлочерепица, Омск, MENTORI"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_august()
