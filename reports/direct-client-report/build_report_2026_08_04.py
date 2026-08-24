from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mentori/mentori-crm")
OUT_DIR = ROOT / "reports/direct-client-report"
OUT = OUT_DIR / "Отчет_Яндекс_Директ_Спектр_Металла_29_июля-4_августа_2026.docx"
ASSET = Path("/Users/mentori/Спектр металл/assets/direct")
MENTORI_LOGO = Path("/Users/mentori/mentra-site/assets/mentori-technologies-wordmark-transparent-light.png")

RED = "C82020"
DARK = "17191F"
CHARCOAL = "292C33"
MID = "626773"
LIGHT = "F3F4F6"
PALE_RED = "FBEAEA"
PALE_GREEN = "EAF5EF"
GREEN = "2E7D50"
WHITE = "FFFFFF"
LINE = "D9DCE2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches, indent_dxa=120):
    widths = [int(v * 1440) for v in widths_inches]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=10.3, bold=False, color=DARK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, before=0, after=6, line=1.12, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep


def add_text(doc, text="", size=10.3, bold=False, color=DARK, after=6,
             before=0, align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, keep=keep)
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=3, line=1.08)
    set_font(p.add_run(text), size=9.8, color=DARK)
    return p


def create_numbering(doc, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    lvl.append(start_node)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text_node = OxmlElement("w:lvlText")
    text_node.set(qn("w:val"), "%1.")
    lvl.append(text_node)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, title, body, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    style_paragraph(p, after=5, line=1.12)
    set_font(p.add_run(title + ". "), size=10.1, bold=True, color=DARK)
    set_font(p.add_run(body), size=10.1, color=DARK)
    return p


def add_callout(doc, label, text, fill=LIGHT, label_color=RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell,
                    top={"val": "single", "sz": 4, "color": fill},
                    bottom={"val": "single", "sz": 4, "color": fill},
                    left={"val": "single", "sz": 22, "color": label_color},
                    right={"val": "single", "sz": 4, "color": fill})
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.15)
    set_font(p.add_run(label + "  "), size=10.2, bold=True, color=label_color)
    set_font(p.add_run(text), size=10.2, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_geometry(table, [6.5 / len(metrics)] * len(metrics))
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, DARK)
        set_cell_border(cell,
                        top={"val": "single", "sz": 4, "color": DARK},
                        bottom={"val": "single", "sz": 4, "color": DARK},
                        left={"val": "single", "sz": 4, "color": CHARCOAL},
                        right={"val": "single", "sz": 4, "color": CHARCOAL})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=1)
        set_font(p.add_run(value), size=18, bold=True, color=WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(label), size=8.4, color="D4D6DB")
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_section_gap(doc, points=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 0.7


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def make_campaign_table(doc):
    data = [
        ["Кампания", "Показы", "Переходы", "CTR", "Расход", "Целевые действия"],
        ["Профнастил", "1 860", "217", "11,67%", "3 057,59 ₽", "3"],
        ["Металлочерепица", "632", "47", "7,44%", "988,63 ₽", "4"],
        ["Итого", "2 492", "264", "10,59%", "4 046,22 ₽", "7"],
    ]
    table = doc.add_table(rows=len(data), cols=6)
    widths = [1.42, .68, .90, .62, 1.10, 1.78]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, DARK)
                color, bold = WHITE, True
            elif r_idx == len(data) - 1:
                set_cell_shading(cell, LIGHT)
                color, bold = DARK, True
            else:
                set_cell_shading(cell, WHITE)
                color, bold = DARK, c_idx == 0
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 5) else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.7, bold=bold, color=color)
    return table


def add_conversion_table(doc):
    rows = [
        ["Клик по телефону", "5 - посетитель открыл набор номера"],
        ["Переход в WhatsApp", "2 - посетитель открыл диалог с компанией"],
        ["Итого в Директе", "7 контактных действий, а не семь оплаченных заказов"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2.05, 4.45])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, LIGHT if r_idx == len(rows) - 1 else WHITE)
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.8,
                     bold=(r_idx == len(rows) - 1 or c_idx == 0), color=DARK)
    return table


def add_traffic_source_table(doc):
    rows = [
        ["Источник", "Переходы", "Расход", "Доля расхода"],
        ["Автотаргетинг", "208", "3 392,09 ₽", "83,8%"],
        ["Ключевые фразы", "56", "654,13 ₽", "16,2%"],
        ["Итого", "264", "4 046,22 ₽", "100%"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [2.35, 1.15, 1.55, 1.45])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx == len(rows) - 1 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.9,
                     bold=(r_idx == 0 or r_idx == len(rows) - 1 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_bad_queries_table(doc):
    rows = [
        ["Тип нецелевого спроса", "Примеры", "Потери"],
        ["Другие города", "Ишим, Новокузнецк, Москва, Челябинск", "отсечение географии"],
        ["Конкуренты и адреса", "Сайдинг Инвест, Основа Строй, 33 Северная, Ю3", "исключение брендов"],
        ["Другие товары", "краска для пола, рифлёный лист, печать, резиновый", "исключение тематики"],
        ["Низкая готовность", "б/у, справочные и информационные запросы", "перенос в SEO"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_geometry(table, [1.62, 3.25, 1.63])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.5, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_seo_table(doc):
    rows = [
        ["№", "Материал", "Спрос за 30 дней*", "Зачем нужен"],
        ["1", "Цены на профнастил: лист, м² и длина 6 м", "1 597 / 696 с Омском", "Коммерческий спрос"],
        ["2", "Какой профнастил выбрать для крыши", "931 / 273 с Омском", "Подводит к кровельным профилям"],
        ["3", "Какой профнастил выбрать для забора", "684 / 209 с Омском", "Подводит к С8, С10, МП-20"],
        ["4", "Размеры: рабочая ширина, длина, толщина", "305 / 542 / 148", "Снимает вопросы перед расчётом"],
        ["5", "Цвета профнастила RAL с примерами", "247", "Помогает выбрать цвет и оставить заявку"],
        ["6", "Виды профнастила: С8-НС44", "63 + профильные запросы", "Сравнение ассортимента"],
        ["7", "Калькулятор профнастила: как рассчитать", "35 + реальные запросы", "Высокая практическая ценность"],
        ["8", "Размеры и рабочая ширина металлочерепицы", "139 / 77", "Поддержка второй кампании"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [.62, 2.82, 1.42, 1.64])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, DARK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.0)
            set_font(p.add_run(value), size=7.9, bold=(r_idx == 0 or c_idx == 0),
                     color=WHITE if r_idx == 0 else DARK)
    return table


def add_image_grid(doc):
    items = [
        (ASSET / "profnastil-fence-16x9.png", "Готовый забор"),
        (ASSET / "profnastil-roof-16x9.png", "Готовая кровля"),
        (ASSET / "profnastil-colors-4x3.png", "Профили и цвета"),
        (ASSET / "metallocherepitsa-fan-1x1.png", "Линейка цветов"),
        (ASSET / "metallocherepitsa-house-4x3.png", "Кровля, графит"),
        (ASSET / "metallocherepitsa-roof-16x9.png", "Кровля, бордо"),
    ]
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [2.166, 2.167, 2.167])
    for idx, (path, caption) in enumerate(items):
        cell = table.cell(idx // 3, idx % 3)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, WHITE)
        set_cell_border(cell,
                        top={"val": "single", "sz": 7, "color": LINE},
                        bottom={"val": "single", "sz": 7, "color": LINE},
                        left={"val": "single", "sz": 7, "color": LINE},
                        right={"val": "single", "sz": 7, "color": LINE})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=2)
        if path.exists():
            shape = p.add_run().add_picture(str(path), width=Inches(1.82), height=Inches(1.05))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(caption), size=8.3, bold=True, color=DARK)
    return table


def add_budget_table(doc):
    rows = [
        ["Этап", "Бюджет", "Как распределяем", "Цель этапа"],
        ["1. Контрольный тест", "10 000 ₽", "Уже потрачено - 4 046 ₽\nОстаток теста - 5 954 ₽", "Дочищаем запросы и сравниваем цену качественного обращения"],
        ["2. Масштабирование", "20 000 ₽", "Доли определяем по подтверждённым обращениям, а не только по кликам", "Усиливаем лучшие товарные группы без автоматического роста бюджета"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1.52, .78, 2.0, 2.20])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "D7D9DE" if r_idx == 0 and c_idx == 3 else (DARK if r_idx == 0 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 5, "color": LINE},
                            bottom={"val": "single", "sz": 5, "color": LINE},
                            left={"val": "single", "sz": 5, "color": LINE},
                            right={"val": "single", "sz": 5, "color": LINE})
            p = cell.paragraphs[0]
            if r_idx == 0 and c_idx == 3:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.08)
            text_color = DARK if r_idx == 0 and c_idx == 3 else (WHITE if r_idx == 0 else DARK)
            set_font(p.add_run(value), size=8.5, bold=(r_idx == 0 or c_idx in (0, 1)), color=text_color)
    return table


def add_header_footer(doc):
    for section in doc.sections:
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(hp, after=0)
        set_font(hp.add_run("СПЕКТР МЕТАЛЛА  |  ЯНДЕКС.ДИРЕКТ  |  MENTORI"), size=7.8, bold=True, color=MID)
        fp = section.footer.paragraphs[0]
        style_paragraph(fp, after=0)
        set_font(fp.add_run("Срез на 4 августа 2026  •  metallomsk.ru  •  mentori.tech"), size=7.8, color=MID)
        fp.add_run(" " * 8)
        run = fp.add_run("Страница ")
        set_font(run, size=7.8, color=MID)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, (size, color, before, after) in enumerate(((18, RED, 16, 8), (13.5, DARK, 11, 5), (11.3, CHARCOAL, 8, 4)), start=1):
        style = styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.8)

    # Page 1: cover and executive summary.
    add_text(doc, "СПЕКТР МЕТАЛЛА", size=10, bold=True, color=RED, after=28)
    add_text(doc, "Результаты рекламы\nи план SEO-продвижения", size=28, bold=True, color=DARK, after=8, keep=True)
    add_text(doc, "Профнастил и металлочерепица  •  Омск и Омская область", size=12.8, bold=True, color=MID, after=3)
    add_text(doc, "Срез за 29 июля - 4 августа 2026 года", size=10, color=MID, after=3)
    add_text(doc, "Подготовлено MENTORI  •  mentori.tech", size=9.4, bold=True, color=RED, after=22)
    add_metric_strip(doc, [
        ("2 492", "показа"),
        ("264", "перехода"),
        ("4 046 ₽", "расход"),
        ("7", "целевых действий"),
    ])
    add_text(doc, "Главный вывод", size=14, bold=True, color=DARK, before=18, after=8)
    add_callout(doc, "Результат", "Обе кампании приводят заинтересованных посетителей. Металлочерепица пока даёт контактные действия дешевле, а профнастил обеспечивает основной объём трафика и самый высокий CTR.", fill=PALE_GREEN, label_color=GREEN)
    add_bullet(doc, "Средняя стоимость перехода снизилась до 15,33 ₽ при установленном пределе 70 ₽.")
    add_bullet(doc, "Зафиксировано 5 кликов по телефону и 2 перехода в WhatsApp.")
    add_bullet(doc, "Нецелевые переходы обнаружены и собраны в отдельный список для дополнительной очистки.")
    add_bullet(doc, "Информационные запросы вынесены в SEO-план, чтобы получать этот спрос без оплаты каждого клика.")
    add_text(doc, "План бюджета", size=14, bold=True, color=DARK, before=12, after=7)
    add_callout(doc, "30 000 ₽", "10 000 ₽ - контролируемый тест и очистка трафика. Оставшиеся 20 000 ₽ - усиление направлений, которые подтверждаются качественными обращениями.", fill=PALE_RED, label_color=RED)

    # Page 2: campaign performance.
    add_page_break(doc)
    section_one_numbers = create_numbering(doc)
    add_heading(doc, "1. Результаты рекламных кампаний", 1)
    add_text(doc, "Целевое действие - попытка связаться с компанией: клик по телефону или переход в мессенджер. Это важный шаг к заявке, но ещё не означает состоявшуюся продажу.", size=9.2, color=MID, after=10)
    make_campaign_table(doc)
    add_heading(doc, "Что показывают цифры", 2)
    add_number(doc, "Профнастил даёт основной объём", "217 переходов, CTR 11,67% и средняя стоимость перехода 14,09 ₽. Кампания собрала 3 контактных действия.", section_one_numbers)
    add_number(doc, "Металлочерепица эффективнее по контакту", "47 переходов привели к 4 контактным действиям. Ориентировочная стоимость одного действия - 247,16 ₽ против 1 019,20 ₽ у профнастила.", section_one_numbers)
    add_number(doc, "Общий CTR высокий", "10,59% посетителей, увидевших рекламу, перешли на сайт. Средняя стоимость перехода по двум кампаниям - 15,33 ₽.", section_one_numbers)
    add_heading(doc, "Сравнение эффективности", 2)
    winners = [
        ("Профнастил", "CPC 14,09 ₽  •  3 действия  •  ориентировочный CPA 1 019,20 ₽"),
        ("Металлочерепица", "CPC 21,03 ₽  •  4 действия  •  ориентировочный CPA 247,16 ₽"),
        ("Итого", "7 контактных действий  •  ориентировочный CPA 578,03 ₽"),
    ]
    table = doc.add_table(rows=len(winners), cols=2)
    set_table_geometry(table, [2.05, 4.45])
    for idx, (name, detail) in enumerate(winners):
        for col, value in enumerate((name, detail)):
            cell = table.cell(idx, col)
            set_cell_shading(cell, LIGHT if idx % 2 == 0 else WHITE)
            set_cell_border(cell,
                            top={"val": "single", "sz": 3, "color": LINE},
                            bottom={"val": "single", "sz": 3, "color": LINE},
                            left={"val": "single", "sz": 3, "color": LINE},
                            right={"val": "single", "sz": 3, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.9, bold=(col == 0 or idx == 2), color=DARK)
    add_callout(doc, "Важно", "В отчёте не называем семь действий семью заявками. Отдел продаж должен подтверждать, какие звонки и сообщения стали реальными обращениями и заказами.", fill=LIGHT, label_color=MID)

    # Page 3: conversions and targeting.
    add_page_break(doc)
    add_heading(doc, "2. Контакты и источники трафика", 1)
    add_text(doc, "Директ учитывает действия посетителей на сайте. Для оценки бизнеса эти данные нужно сопоставлять с письмами, звонками и фактическими диалогами менеджера.", after=9)
    add_text(doc, "Какие действия зафиксированы", size=11.3, bold=True, color=DARK, after=5)
    add_conversion_table(doc)
    add_callout(doc, "Отдельный факт", "Подтверждена реальная заявка, пришедшая на почту с рекламы. Она не прибавлена к семи действиям, чтобы исключить возможное задвоение.", fill=PALE_GREEN, label_color=GREEN)
    add_heading(doc, "Роль автотаргетинга", 2)
    add_traffic_source_table(doc)
    add_bullet(doc, "Автотаргетинг дал 208 из 264 переходов и использовал 83,8% расходов.")
    add_bullet(doc, "Полностью отключать его сейчас невыгодно: он обеспечивает основной объём и уже приводит контактные действия.")
    add_bullet(doc, "Правильная тактика - регулярно чистить поисковые запросы и ограничивать явно нецелевые темы.")
    add_callout(doc, "Статус", "Кампании работают только в поиске Яндекса, по Омску и Омской области, ежедневно с 08:00 до 21:00 по омскому времени.", fill=LIGHT, label_color=MID)

    # Page 4: search-query cleanup.
    add_page_break(doc)
    section_queries = create_numbering(doc)
    add_heading(doc, "3. Качество поисковых запросов", 1)
    add_callout(doc, "Обнаружено", "14 явно нецелевых формулировок: 17 переходов, 350,44 ₽ расходов и 0 целевых действий. Это 8,7% всех расходов за период.", fill=PALE_RED, label_color=RED)
    add_bad_queries_table(doc)
    add_heading(doc, "Что делаем с этим трафиком", 2)
    add_number(doc, "Расширяем минус-фразы", "Добавляем фамилии и названия компаний, другие города, адресные запросы и неподходящие товары.", section_queries)
    add_number(doc, "Не блокируем полезный спрос", "Запрос «металлический штакетник Омск» уже дал контактное действие. Его лучше вынести в отдельную посадочную страницу и рекламную группу.", section_queries)
    add_number(doc, "Информационный спрос переносим в SEO", "39 справочных формулировок дали 44 перехода, 321,63 ₽ расходов и 0 целевых действий. На них лучше отвечать статьями, а не оплачивать каждый переход.", section_queries)
    add_heading(doc, "Примеры запросов для исключения", 2)
    add_bullet(doc, "«стоимость профлиста в Новокузнецке», «профнастил оцинковка Ишим»;")
    add_bullet(doc, "«ИП Игнатенко Анатолий Владимирович», «Сайдинг Инвест», «Основа Строй»;")
    add_bullet(doc, "«краска для пола», «профлист печать», «профнастил б/у», «лист рифлёный».")
    add_text(doc, "Часть этих исключений уже внесена. Дополнительный список подготовлен для следующего обновления кампаний.", size=9.4, bold=True, color=RED, before=7, after=0)

    # Page 5: SEO plan.
    add_page_break(doc)
    add_heading(doc, "4. Темы для SEO-продвижения", 1)
    add_text(doc, "Поисковые запросы из рекламы показывают, какие ответы нужны покупателю до обращения. Эти темы можно закрыть полезными страницами и постепенно получать переходы без оплаты каждого клика.", after=9)
    add_seo_table(doc)
    add_text(doc, "* Данные Wordstat по Омской области за последние 30 дней. Это количество запросов с фразой, а не число уникальных покупателей.", size=8.5, italic=True, color=MID, before=5, after=8)
    add_heading(doc, "Как должна работать каждая статья", 2)
    add_bullet(doc, "Короткий ответ на вопрос покупателя в начале страницы.")
    add_bullet(doc, "Таблица размеров, покрытий, цен или вариантов применения.")
    add_bullet(doc, "Ссылки на нужную категорию каталога, калькулятор и контакты.")
    add_bullet(doc, "Кнопка расчёта и понятный призыв оставить заявку.")
    add_callout(doc, "Приоритет", "Сначала публикуем материалы про цену, крышу, забор и размеры. Они объединяют самый большой спрос и напрямую связаны с продажей профнастила.", fill=PALE_GREEN, label_color=GREEN)

    # Page 6: messages and visual tests.
    add_page_break(doc)
    add_heading(doc, "5. Что тестируем в объявлениях", 1)
    add_text(doc, "Рекламные группы разделены по задаче покупателя. Это позволяет сравнить не только товары, но и аргументы, формулировки и изображения.", after=9)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3.19, 3.31])
    blocks = [
        ("ПРОФНАСТИЛ", [
            "Купить профнастил в Омске",
            "Профнастил для забора",
            "Профлист для крыши",
            "Профили С8-НС44",
            "Цвета RAL и размеры",
        ]),
        ("МЕТАЛЛОЧЕРЕПИЦА", [
            "Цена и покупка в Омске",
            "Металлочерепица для крыши",
            "Кровельный комплект",
            "Супер-Монтеррей",
            "Размеры и выбор цвета",
        ]),
    ]
    for col, (title, lines) in enumerate(blocks):
        cell = table.cell(0, col)
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell,
                        top={"val": "single", "sz": 5, "color": LINE},
                        bottom={"val": "single", "sz": 5, "color": LINE},
                        left={"val": "single", "sz": 5, "color": LINE},
                        right={"val": "single", "sz": 5, "color": LINE})
        p = cell.paragraphs[0]
        style_paragraph(p, after=5)
        set_font(p.add_run(title), size=10.2, bold=True, color=RED)
        for line in lines:
            p = cell.add_paragraph(style="List Bullet")
            style_paragraph(p, after=2, line=1.05)
            set_font(p.add_run(line), size=8.8, color=DARK)
    add_heading(doc, "Какие аргументы и фотографии сравниваем", 2)
    add_bullet(doc, "Цена, изготовление по размерам, выбор RAL, доставка по Омску и области.")
    add_bullet(doc, "Готовый объект против наглядной фотографии отдельного профиля или цветовой линейки.")
    add_image_grid(doc)
    add_text(doc, "Следующий вывод по визуалам делаем не по вкусу, а по фактической цене качественного обращения после накопления данных.", size=9.1, color=MID, before=7, after=0)

    # Page 7: budget and next steps.
    add_page_break(doc)
    section_next = create_numbering(doc)
    add_heading(doc, "6. Бюджет и следующий этап", 1)
    add_text(doc, "Общий план остаётся прежним: сначала ограниченный тест, затем масштабирование только подтверждённых связок.", after=10)
    add_budget_table(doc)
    add_callout(doc, "Текущий остаток", "Из тестовых 10 000 ₽ использовано 4 046,22 ₽. До контрольной точки остаётся 5 953,78 ₽.", fill=PALE_RED, label_color=RED)
    add_heading(doc, "Что делаем дальше", 2)
    add_number(doc, "Дочищаем запросы", "вносим новые минус-фразы и ежедневно проверяем, куда уходит бюджет автотаргетинга.", section_next)
    add_number(doc, "Продолжаем обе кампании", "профнастил сохраняет сильный спрос, металлочерепица пока показывает более низкую стоимость контактного действия.", section_next)
    add_number(doc, "Сверяем с продажами", "менеджер помечает качество каждого звонка, сообщения и письма, чтобы оптимизация шла по реальным обращениям.", section_next)
    add_number(doc, "Запускаем SEO-публикации", "начинаем с цены, выбора для крыши и забора, затем закрываем размеры, цвета и калькулятор.", section_next)
    add_number(doc, "После тестовых 10 000 ₽", "распределяем оставшиеся 20 000 ₽ по цене качественной заявки и согласованным приоритетам бизнеса.", section_next)
    add_heading(doc, "Что нужно подтвердить у клиента", 2)
    add_bullet(doc, "Какие профили, покрытия и цвета сейчас важнее по наличию и маржинальности.")
    add_bullet(doc, "Какие обращения считаются качественными и какие завершились продажей.")
    add_bullet(doc, "Нужно ли следующим отдельным направлением запускать евроштакетник.")
    add_callout(doc, "Итог", "Реклама уже даёт измеримые контактные действия. Ближайшая задача - уменьшить долю нецелевого спроса и связать статистику Директа с реальными продажами.", fill=PALE_GREEN, label_color=GREEN)
    add_text(doc, "Отчёт подготовлен MENTORI  •  mentori.tech", size=10.2, bold=True, color=RED, before=9, after=0)

    add_header_footer(doc)
    doc.core_properties.title = "Результаты Яндекс.Директа и SEO-план - Спектр Металла"
    doc.core_properties.subject = "Профнастил и металлочерепица, 29 июля - 4 августа 2026"
    doc.core_properties.author = "MENTORI"
    doc.core_properties.keywords = "Яндекс.Директ, SEO, профнастил, металлочерепица, Омск, отчёт, MENTORI"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
