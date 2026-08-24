from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mentori/mentori-crm")
OUT_DIR = ROOT / "reports/direct-client-report"
OUT = OUT_DIR / "Отчет_Яндекс_Директ_Спектр_Металла_30-31_июля_2026.docx"
ASSET = Path("/Users/mentori/Спектр металл/assets/direct")

RED = "C82020"
DARK = "17191F"
CHARCOAL = "292C33"
MID = "626773"
LIGHT = "F3F4F6"
PALE_RED = "FBEAEA"
PALE_GREEN = "EAF5EF"
GREEN = "2E7D50"
WHITE = "FFFFFF"
LINE = "D9DCE2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches, indent_dxa=120):
    widths = [int(v * 1440) for v in widths_inches]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=10.3, bold=False, color=DARK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, before=0, after=6, line=1.12, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep


def add_text(doc, text="", size=10.3, bold=False, color=DARK, after=6,
             before=0, align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, keep=keep)
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=3, line=1.08)
    set_font(p.add_run(text), size=9.8, color=DARK)
    return p


def create_numbering(doc, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    lvl.append(start_node)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text_node = OxmlElement("w:lvlText")
    text_node.set(qn("w:val"), "%1.")
    lvl.append(text_node)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, title, body, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    style_paragraph(p, after=5, line=1.12)
    set_font(p.add_run(title + ". "), size=10.1, bold=True, color=DARK)
    set_font(p.add_run(body), size=10.1, color=DARK)
    return p


def add_callout(doc, label, text, fill=LIGHT, label_color=RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell,
                    top={"val": "single", "sz": 4, "color": fill},
                    bottom={"val": "single", "sz": 4, "color": fill},
                    left={"val": "single", "sz": 22, "color": label_color},
                    right={"val": "single", "sz": 4, "color": fill})
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.15)
    set_font(p.add_run(label + "  "), size=10.2, bold=True, color=label_color)
    set_font(p.add_run(text), size=10.2, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_geometry(table, [6.5 / len(metrics)] * len(metrics))
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, DARK)
        set_cell_border(cell,
                        top={"val": "single", "sz": 4, "color": DARK},
                        bottom={"val": "single", "sz": 4, "color": DARK},
                        left={"val": "single", "sz": 4, "color": CHARCOAL},
                        right={"val": "single", "sz": 4, "color": CHARCOAL})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=1)
        set_font(p.add_run(value), size=18, bold=True, color=WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(label), size=8.4, color="D4D6DB")
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_section_gap(doc, points=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 0.7


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def make_campaign_table(doc):
    data = [
        ["Кампания", "Показы", "Переходы", "CTR", "Расход", "Результат"],
        ["Профнастил", "231", "29", "12,55%", "1 114,79 ₽", "1 заявка"],
        ["Металлочерепица", "166", "7", "4,22%", "275,40 ₽", "Данные набираются"],
        ["Итого", "397", "36", "9,07%", "1 390,19 ₽", "1 заявка"],
    ]
    table = doc.add_table(rows=len(data), cols=6)
    widths = [1.52, .68, .82, .62, 1.05, 1.81]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, DARK)
                color, bold = WHITE, True
            elif r_idx == len(data) - 1:
                set_cell_shading(cell, LIGHT)
                color, bold = DARK, True
            else:
                set_cell_shading(cell, WHITE)
                color, bold = DARK, c_idx == 0
            set_cell_border(cell,
                            top={"val": "single", "sz": 4, "color": LINE},
                            bottom={"val": "single", "sz": 4, "color": LINE},
                            left={"val": "single", "sz": 4, "color": LINE},
                            right={"val": "single", "sz": 4, "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 5) else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.7, bold=bold, color=color)
    return table


def add_image_grid(doc):
    items = [
        (ASSET / "profnastil-fence-16x9.png", "Готовый забор"),
        (ASSET / "profnastil-roof-16x9.png", "Готовая кровля"),
        (ASSET / "profnastil-colors-4x3.png", "Профили и цвета"),
        (ASSET / "metallocherepitsa-fan-1x1.png", "Линейка цветов"),
        (ASSET / "metallocherepitsa-house-4x3.png", "Кровля, графит"),
        (ASSET / "metallocherepitsa-roof-16x9.png", "Кровля, бордо"),
    ]
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [2.166, 2.167, 2.167])
    for idx, (path, caption) in enumerate(items):
        cell = table.cell(idx // 3, idx % 3)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, WHITE)
        set_cell_border(cell,
                        top={"val": "single", "sz": 7, "color": LINE},
                        bottom={"val": "single", "sz": 7, "color": LINE},
                        left={"val": "single", "sz": 7, "color": LINE},
                        right={"val": "single", "sz": 7, "color": LINE})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=2)
        if path.exists():
            shape = p.add_run().add_picture(str(path), width=Inches(1.82), height=Inches(1.05))
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p2, after=0, line=1.0)
        set_font(p2.add_run(caption), size=8.3, bold=True, color=DARK)
    return table


def add_budget_table(doc):
    rows = [
        ["Этап", "Бюджет", "Как распределяем", "Цель этапа"],
        ["1. Контрольный тест", "10 000 ₽", "Профнастил — 5 500 ₽\nМеталлочерепица — 3 500 ₽\nРезерв — 1 000 ₽", "Сравниваем направления, тексты, фотографии и реальные поисковые запросы"],
        ["2. Масштабирование", "20 000 ₽", "Предварительно: 12 000 ₽ на профнастил и 8 000 ₽ на металлочерепицу", "Параллельно усиливаем только те группы, которые дают качественные обращения"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1.52, .78, 2.0, 2.20])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "D7D9DE" if r_idx == 0 and c_idx == 3 else (DARK if r_idx == 0 else WHITE))
            set_cell_border(cell,
                            top={"val": "single", "sz": 5, "color": LINE},
                            bottom={"val": "single", "sz": 5, "color": LINE},
                            left={"val": "single", "sz": 5, "color": LINE},
                            right={"val": "single", "sz": 5, "color": LINE})
            p = cell.paragraphs[0]
            if r_idx == 0 and c_idx == 3:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.08)
            text_color = DARK if r_idx == 0 and c_idx == 3 else (WHITE if r_idx == 0 else DARK)
            set_font(p.add_run(value), size=8.5, bold=(r_idx == 0 or c_idx in (0, 1)), color=text_color)
    return table


def add_header_footer(doc):
    for section in doc.sections:
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(hp, after=0)
        set_font(hp.add_run("СПЕКТР МЕТАЛЛА  |  ЯНДЕКС.ДИРЕКТ"), size=7.8, bold=True, color=MID)
        fp = section.footer.paragraphs[0]
        style_paragraph(fp, after=0)
        set_font(fp.add_run("Отчёт за 30–31 июля 2026  •  metallomsk.ru"), size=7.8, color=MID)
        fp.add_run(" " * 8)
        run = fp.add_run("Страница ")
        set_font(run, size=7.8, color=MID)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, (size, color, before, after) in enumerate(((18, RED, 16, 8), (13.5, DARK, 11, 5), (11.3, CHARCOAL, 8, 4)), start=1):
        style = styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.8)

    # Page 1: cover and executive summary.
    add_text(doc, "СПЕКТР МЕТАЛЛА", size=10, bold=True, color=RED, after=34)
    add_text(doc, "Результаты тестового запуска\nЯндекс.Директа", size=29, bold=True, color=DARK, after=8, keep=True)
    add_text(doc, "Профнастил и металлочерепица  •  Омск и Омская область", size=12.8, bold=True, color=MID, after=3)
    add_text(doc, "Отчёт за 30–31 июля 2026 года", size=10, color=MID, after=28)
    add_metric_strip(doc, [
        ("397", "показов"),
        ("36", "переходов"),
        ("1 390 ₽", "расход"),
        ("1", "заявка"),
    ])
    add_text(doc, "Первые выводы", size=14, bold=True, color=DARK, before=20, after=8)
    add_callout(doc, "Главное", "Кампания по профнастилу уже дала первую заявку и показывает высокий интерес аудитории. Кампания по металлочерепице запущена позже и пока находится на этапе накопления данных.", fill=PALE_GREEN, label_color=GREEN)
    add_bullet(doc, "Средняя стоимость перехода по двум кампаниям — 38,62 ₽.")
    add_bullet(doc, "Профнастил привлекает переходы заметно активнее: CTR 12,55% против 4,22% у металлочерепицы.")
    add_bullet(doc, "Двух дней недостаточно для окончательного решения: сейчас мы отсекаем нецелевые запросы и продолжаем контролируемый тест.")
    add_text(doc, "План на месяц", size=14, bold=True, color=DARK, before=14, after=7)
    add_callout(doc, "30 000 ₽", "10 000 ₽ — проверка гипотез и выбор лучших связок. Затем 20 000 ₽ — параллельное усиление тех направлений, которые подтвердят результат заявками.", fill=PALE_RED, label_color=RED)

    # Page 2: performance.
    add_page_break(doc)
    section_one_numbers = create_numbering(doc)
    add_heading(doc, "1. Что получили за два дня", 1)
    add_text(doc, "Показ — объявление увидели. Переход — человек открыл сайт. CTR — доля людей, которые заинтересовались объявлением и перешли на сайт.", size=9.2, color=MID, after=10)
    make_campaign_table(doc)
    add_heading(doc, "Что видно уже сейчас", 2)
    add_number(doc, "Профнастил — текущий лидер", "29 переходов, высокий интерес и первая зафиксированная заявка. Это направление имеет смысл продолжать тестировать в приоритете.", section_one_numbers)
    add_number(doc, "Металлочерепица — тест продолжается", "7 переходов — пока слишком мало для честного вывода. Лучше всего откликаются объявления для крыши и кровли.", section_one_numbers)
    add_number(doc, "Цена перехода контролируется", "38–39 ₽ в среднем по обеим кампаниям. Это ниже установленного предела 70 ₽ за переход.", section_one_numbers)
    add_section_gap(doc)
    add_heading(doc, "Какие группы показали лучший отклик", 2)
    winners = [
        ("Профнастил: для крыши", "CTR 14,63%  •  6 переходов  •  27,90 ₽ за переход"),
        ("Профнастил: цвета и размеры", "CTR 15,79%  •  6 переходов  •  требуется снизить цену перехода"),
        ("Металлочерепица: для крыши", "CTR 10,53%  •  4 перехода  •  35,25 ₽ за переход"),
        ("Супер-Монтеррей", "CTR 1,72%  •  формулировки и аудиторию нужно дорабатывать"),
    ]
    table = doc.add_table(rows=len(winners), cols=2)
    set_table_geometry(table, [2.35, 4.15])
    for idx, (name, detail) in enumerate(winners):
        for col, value in enumerate((name, detail)):
            cell = table.cell(idx, col)
            set_cell_shading(cell, LIGHT if idx % 2 == 0 else WHITE)
            set_cell_border(cell,
                            top={"val": "single", "sz": 3, "color": LINE},
                            bottom={"val": "single", "sz": 3, "color": LINE},
                            left={"val": "single", "sz": 3, "color": LINE},
                            right={"val": "single", "sz": 3, "color": LINE})
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            set_font(p.add_run(value), size=8.9, bold=(col == 0), color=DARK if idx < 3 else RED)
    add_text(doc, "Важно: одна заявка на 36 переходов — положительный первый сигнал, но выбор победителей делаем только после накопления достаточного объёма данных.", size=9.2, italic=True, color=MID, before=8, after=0)

    # Page 3: message and visual tests.
    add_page_break(doc)
    add_heading(doc, "2. Что именно мы тестируем", 1)
    add_text(doc, "Реклама разделена по реальной задаче покупателя. Благодаря этому можно понять не только какой товар интересен, но и какой аргумент приводит человека на сайт.", after=9)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3.19, 3.31])
    blocks = [
        ("ПРОФНАСТИЛ", [
            "Купить профнастил в Омске",
            "Профнастил для забора в Омске",
            "Профлист для крыши в Омске",
            "Профнастил С8—НС44 в Омске",
            "Цветной профнастил в Омске",
        ]),
        ("МЕТАЛЛОЧЕРЕПИЦА", [
            "Металлочерепица в Омске от 779 ₽/м²",
            "Купить металлочерепицу в Омске",
            "Металлочерепица для крыши в Омске",
            "Кровля из металлочерепицы в Омске",
            "Супер-Монтеррей в Омске от 779 ₽/м²",
        ]),
    ]
    for col, (title, lines) in enumerate(blocks):
        cell = table.cell(0, col)
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell,
                        top={"val": "single", "sz": 5, "color": LINE},
                        bottom={"val": "single", "sz": 5, "color": LINE},
                        left={"val": "single", "sz": 5, "color": LINE},
                        right={"val": "single", "sz": 5, "color": LINE})
        p = cell.paragraphs[0]
        style_paragraph(p, after=5)
        set_font(p.add_run(title), size=10.2, bold=True, color=RED)
        for line in lines:
            p = cell.add_paragraph(style="List Bullet")
            style_paragraph(p, after=2, line=1.05)
            set_font(p.add_run(line), size=8.8, color=DARK)
    add_heading(doc, "Какие аргументы проверяем", 2)
    add_bullet(doc, "Цена от 779 ₽/м² и возможность получить точный расчёт.")
    add_bullet(doc, "Подбор профиля, толщины, покрытия и цветов RAL.")
    add_bullet(doc, "Изготовление листов нужной длины и доставка по Омску и области.")
    add_bullet(doc, "Разные сценарии применения: крыша, забор, фасад, конкретные профили.")
    add_heading(doc, "Какие изображения проверяем", 2)
    add_image_grid(doc)
    add_text(doc, "Сравниваем две логики: понятный вид готового объекта и наглядная демонстрация профиля/цветов. Это поможет выбрать фотографии, которые лучше приводят покупателей.", size=9.1, color=MID, before=7, after=0)

    # Page 4: optimizations and approval requests.
    add_page_break(doc)
    section_three_numbers = create_numbering(doc)
    add_heading(doc, "3. Что уже улучшили после первых данных", 1)
    add_number(doc, "Очистили поисковые запросы", "Добавили минус-фразы, чтобы не платить за запросы из других городов, смежные товары и информационный интерес без намерения купить.", section_three_numbers)
    add_number(doc, "Разделили аудиторию по задачам", "Отдельно тестируются покупка и цена, забор, крыша, профили, цвета и размеры.", section_three_numbers)
    add_number(doc, "Усилили объявления", "Добавили разные заголовки и тексты, быстрые ссылки, преимущества, кнопку «Узнать цену» и контактные данные компании.", section_three_numbers)
    add_number(doc, "Обновили изображения", "Кроме товарных фото на белом фоне используются готовый забор, готовая кровля и цветовые подборки.", section_three_numbers)
    add_number(doc, "Настроили рабочее время", "Показы идут по будням с 08:00 до 21:00 по Омску. Ночью и в выходные бюджет не расходуется.", section_three_numbers)
    add_callout(doc, "Статус", "Все активные объявления прошли модерацию. Сейчас тест идёт только в поиске Яндекса по Омску и Омской области — это снижает риск распыления небольшого бюджета.", fill=PALE_GREEN, label_color=GREEN)
    add_heading(doc, "Что просим подтвердить у клиента", 2)
    add_bullet(doc, "Какие товары сейчас важнее продавать: профнастил, металлочерепица или другое направление.")
    add_bullet(doc, "Можно ли оставлять в рекламе цену «от 779 ₽/м²» и какие предложения нужно выделить сильнее.")
    add_bullet(doc, "Какие покрытия, цвета и профили нужно продвигать в первую очередь по наличию и маржинальности.")
    add_bullet(doc, "Какие фотографии точнее соответствуют реальному товару: готовые объекты или отдельные листы/профили.")
    add_bullet(doc, "Есть ли направления, на которые сейчас не нужно тратить рекламный бюджет.")
    add_text(doc, "Если клиент скорректирует приоритеты сейчас, второй этап будет строиться не только по кликам, но и по реальной выгоде бизнеса.", size=9.5, bold=True, color=RED, before=8, after=0)

    # Page 5: budget and next steps.
    add_page_break(doc)
    section_four_numbers = create_numbering(doc)
    add_heading(doc, "4. План бюджета: 10 000 ₽ + 20 000 ₽", 1)
    add_text(doc, "Бюджет не расходуется сразу. Мы идём в два этапа и после каждого перераспределяем деньги в пользу фактических результатов.", after=10)
    add_budget_table(doc)
    add_callout(doc, "Уже потрачено", "1 390,19 ₽ входят в тестовые 10 000 ₽. На продолжение первого этапа остаётся 8 609,81 ₽.", fill=PALE_RED, label_color=RED)
    add_heading(doc, "Как принимаем решение о масштабировании", 2)
    add_number(doc, "Каждые 2–3 дня", "проверяем запросы, стоимость перехода, заявки и качество обращений.", section_four_numbers)
    add_number(doc, "После 7–10 дней или достаточного объёма кликов", "сравниваем направления и отключаем слабые связки.", section_four_numbers)
    add_number(doc, "Оставшиеся 20 000 ₽", "запускаем параллельно на лучшие группы. Предварительное распределение — 12 000 ₽ на профнастил и 8 000 ₽ на металлочерепицу, но доли меняются по цене качественной заявки.", section_four_numbers)
    add_section_gap(doc)
    add_heading(doc, "Что получит клиент", 2)
    add_bullet(doc, "Понятный еженедельный отчёт: потрачено, переходы, заявки, цена обращения.")
    add_bullet(doc, "Список объявлений и фотографий, которые реально работают лучше остальных.")
    add_bullet(doc, "Отдельный вывод по товарным направлениям: что продолжать, что остановить, что протестировать следующим.")
    add_bullet(doc, "Контроль расходов без автоматического увеличения бюджета.")
    add_callout(doc, "Важно", "Прогноз по количеству заявок пока не фиксируем: выборка слишком мала. Цель теста — получить достаточные данные и затем масштабировать результат, а не обещание без статистики.", fill=LIGHT, label_color=MID)
    add_text(doc, "Следующая контрольная точка: после расходования тестовых 10 000 ₽ или раньше, если накопится достаточно качественных заявок.", size=10.2, bold=True, color=DARK, before=10, after=0)

    add_header_footer(doc)
    doc.core_properties.title = "Результаты тестового запуска Яндекс.Директа — Спектр Металла"
    doc.core_properties.subject = "Профнастил и металлочерепица, 30–31 июля 2026"
    doc.core_properties.author = "Спектр Металла"
    doc.core_properties.keywords = "Яндекс.Директ, профнастил, металлочерепица, Омск, отчёт"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
